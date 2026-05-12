"""DOCX parser — uses python-docx if installed."""

from __future__ import annotations

from pathlib import Path

from doc_structure_parser.models import DocumentStructure, Heading, Paragraph, Table


def parse_docx(path: str | Path) -> DocumentStructure:
    """Parse a DOCX file into a :class:`DocumentStructure`."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "python-docx is not installed. Install with `pip install doc-structure-parser[docx]`."
        ) from exc

    path = Path(path)
    doc = Document(str(path))
    blocks = []
    src_index = 0

    # Walk paragraphs in document order.
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name if para.style else "") or ""
        if style.startswith("Heading"):
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 1
            level = max(1, min(level, 6))
            blocks.append(Heading(text=text, level=level, source_index=src_index, metadata={"style": style}))
        else:
            blocks.append(Paragraph(text=text, source_index=src_index, metadata={"style": style}))
        src_index += 1

    # Tables come from doc.tables (already in document order if iterating both).
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if not rows:
            continue
        text_lin = _linearize_table(rows)
        blocks.append(Table(text=text_lin, rows=rows, source_index=src_index))
        src_index += 1

    return DocumentStructure(source=str(path), format="docx", blocks=blocks)


def _linearize_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header = rows[0]
    out: list[str] = []
    for r in rows[1:]:
        out.append(", ".join(f"{header[i] if i < len(header) else 'col_' + str(i)}: {v}" for i, v in enumerate(r)))
    return "\n".join(out) or "(empty table)"
